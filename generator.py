from docx import Document
from docx2pdf import convert
from faker import Faker
from progressbar import ProgressBar
import argparse
import datetime
import time
import random, string
import os, sys


address = ""
prices = []
quantities = []
date = ''
total = 0
taux = 0
heures = 0
heures_supp = 0
sal_brut = 0
charges = 0
total_frais = 0
school = ''
nom_fam = ''
prenom = ''
date_covid = ''
ine = ''
nom_bourse = ''
prenom_bourse = ''
date_f = ''


def main():
    pbar = ProgressBar()
    my_parser = argparse.ArgumentParser(description='Location of the template')
    my_parser.add_argument('Path',
                           metavar='path',
                           type=str,
                           help='the path to the template')

    my_parser.add_argument('Quantity',
                           metavar='qty',
                           type=int,
                           help='the number of file you want')

    args = my_parser.parse_args()
    input_path = args.Path
    qty_file = args.Quantity

    if not os.path.isfile(input_path):
        print('The template specified does not exist')
        sys.exit()

    template_file_path = input_path
    output_file_path = 'output/'

    variables = {
        "${EMPLOEE_NAME}": get_name,
        "${EMPLOEE_PHONE}": get_phone_number,
        "${NOM_FAM}": get_nom_fam,
        "${PRENOM}": get_prenom,
        "${HEURE_DELIV}": get_heure_deliv,
        "${ID_DOC}": get_id_doc,
        "${CLE_CONTROL}": get_cle_control,
        "${SEXE}": get_sexe,
        "${DATE_NAISS}": get_date_naiss,
        "${VILLE_NAISS}": get_ville_naiss,
        "${START_DATE}": get_start_date,
        "${START_DATE_F}": get_start_date_f,
        "${END_DATE}": get_end_date,
        "${ADDRESS}": get_address,
        "${CITY_ADDRESS}": get_city_address,
        "${CLUB}": get_club,
        "${OBJET_REM}": get_objet_rem,
        "${TOTAL_FRAIS}": get_total_frais,
        "${JOB}": get_job,
        "${PRIX}": get_prix,
        "${ITEM}": get_item,
        "${QUANTITY}": get_quantity,
        "${INTER}": get_inter,
        "${TOTAL}": get_total,
        "${COMPANY}": get_company,
        "${HEURES}": get_heures,
        "${HEURES_SUPP}": get_heures_supp,
        "${TAUX}": get_taux_horaire,
        "${SAL_HEURES}": get_sal_heures,
        "${SAL_HEURES_SUPP}": get_sal_heures_supp,
        "${SALAIRE_BRUT}": get_brut_salary,
        "${CHARGES}": get_social_charge,
        "${SALAIRE_NET}": get_net_salary,
        "${SCHOOL}": get_school,
        "${ID_DOSS}": get_id_doss,
        "${DATE_COVID}": get_date_covid,
        "${ID_INTERNET}": get_id_internet,
        "${REF_BIL}": get_ref_billet,
        "${REF_VOL}": get_ref_vol,
        "${HEURE_D}": get_heure_d,
        "${HEURE_A}": get_heure_a,
        "${HEURE_D-1}": get_heure_d_1,
        "${NUM_TICK}": get_num_ticket,
        "${DATE_VOL}": get_date_vol,
        "${PAYS}": get_pays,
        "${INE}": get_ine,
        "${REF_BOURSE}": get_ref_bourse,
        "${NOM_BOURSE}": get_nom_bourse,
        "${PRENOM_BOURSE}": get_prenom_bourse,
        "${EMAIL}": get_email,

    }

    pbar = ProgressBar(maxval=qty_file).start()
    for j in range(0, qty_file):
        global date_covid
        date_covid = ''

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)
        template_document.save(
            output_file_path + template_file_path.split('/')[1].split('.')[0] + str(time.time()) + ".docx")
        pbar.update(j+1)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs

        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value())


def get_id_doss():
    res = ""
    for _ in range(0, 10):
        res += str(random.choice(string.ascii_uppercase + string.digits))
    return res


def get_date_covid():
    global date_covid
    fake = Faker()
    if date_covid == '':
        date_covid = fake.date_between('-1y')
    return str(date_covid)


def get_id_internet():
    res = "P-000"
    for _ in range(0, 6):
        res += str(random.randint(0, 9))
    return str(res)


def get_total_frais():
    global total_frais
    if total_frais != 0:
        temp = total_frais
        total_frais = 0
        return str(temp)
    else:
        total_frais = round(random.uniform(10.0, 100.0), 2)
        return str(total_frais)


def get_objet_rem():
    return random.choice(list(open('rembours.txt')))


def get_club():
    return random.choice(list(open('clubs.txt')))


def get_sexe():
    return random.choice(["Masculin", "Féminin"])


def get_heure_deliv():
    fake = Faker()
    return str(fake.time())


def get_date_naiss():
    fake = Faker()
    return str(fake.date_between('-60y', '-20y'))


def get_ville_naiss():
    fake = Faker('fr_FR')
    return fake.city()


def get_id_doc():
    res = ''
    for _ in range(0, 13):
        res += random.choice(['0', '1', '2', '3', '4', '5', '6', '7', '8', '9', 'A', 'B', 'C', 'D', 'E', 'F'])
    return res


def get_cle_control():
    res = ''
    for _ in range(0, 8):
        res += random.choice(['0', '1', '2', '3', '4', '5', '6', '7', '8', '9', 'A', 'B', 'C', 'D', 'E', 'F'])
    return res


def get_nom_fam():
    global nom_fam
    fake = Faker('fr_FR')
    if nom_fam != '':
        temp = nom_fam
        nom_fam = ''
        return temp
    else:
        nom_fam = fake.last_name()
        return nom_fam


def get_prenom():
    global prenom
    fake = Faker('fr_FR')
    if prenom != '':
        temp = prenom
        prenom = ''
        return temp
    else:
        prenom = fake.first_name_male()
        return prenom


def get_phone_number():
    nb = '+33'
    for i in range(1, 10):
        nb = nb + str(random.randint(0, 9))
    return nb


def get_prix():
    price = random.randrange(10, 150, 10)
    global prices
    prices.append(price)
    return str(price)


def get_item():
    return random.choice(list(open('construction.txt')))


def get_school():
    global school
    if school != '':
        temp = school
        school = ''
        return temp
    else:
        school = random.choice(list(open('ecole.txt')))
        return school


def get_quantity():
    q = random.choice([1, 2, 5, 10, 15, 20, 25])
    global quantities
    quantities.append(q)
    return str(q)


def get_company():
    fake = Faker()
    return fake.company()


def get_job():
    fake = Faker('fr_FR')
    return fake.job()


def get_start_date():
    fake = Faker()
    global date
    d = ''
    if date == '':
        d = fake.date_between('-5y', 'today')
        date = d
    else:
        d = date
        date = ''
    return str(d)

def get_start_date_f():
    fake = Faker()
    global date_f
    date_f = fake.date_between('-5y', 'today')
    return str(date_f)


def get_end_date():
    global date_f
    fake = Faker()
    if date_f == '':
        date_f = fake.date_between('-5y', 'today')
    delta = datetime.timedelta(days=random.randint(5, 90))
    end_date = date_f + delta
    return str(end_date)


def get_address():
    global address
    fake = Faker('fr_FR')
    add = fake.address()
    address = add
    return add


def get_city_address():
    global address
    c = address.split(' ')[-1]
    return c


def get_name():
    fake = Faker('fr_FR')
    return fake.name()


def get_inter():
    global prices
    global quantities
    res = prices.pop(0) * quantities.pop(0)
    global total
    total = total + res
    return str(res)


def get_total():
    global total
    res = total
    total = 0
    return str(res)


def get_heures():
    global heures
    h = random.randint(120, 200)
    heures = h
    return str(h)


def get_heures_supp():
    global heures_supp
    h = random.choice([0, 0, 0, 0, 0, 5, 10])
    heures_supp = h
    return str(h)


def get_taux_horaire():
    global taux
    t = round(random.uniform(10.03, 15.0), 2)
    taux = t
    return str(t)


def get_sal_heures():
    global heures
    global taux
    return str(heures * taux)


def get_sal_heures_supp():
    global heures_supp
    global taux
    return str(heures_supp * (taux * 1.25))


def get_brut_salary():
    global sal_brut
    global heures
    global heures_supp
    global taux
    s = round((heures * taux + heures_supp * (taux * 1.25)), 2)
    sal_brut = s
    return str(s)


def get_social_charge():
    global sal_brut
    global charges
    c = round(sal_brut * 0.22, 2)
    charges = c
    return str(c)


def get_net_salary():
    global sal_brut
    global charges
    s = round(sal_brut - charges, 2)
    return str(s)

def get_date_vol():
    res = ''
    res += str(random.randint(1,31))
    res += random.choice(['JAN','FEV','MAR','AVR','MAI','JUN','JUI','AOU','SEP','OCT','NOV','DEC'])
    return str(res)

def get_num_ticket():
    res = ''
    for i in range(4):
        for j in range(3):
            res += str(random.randint(0,9))
        res += ' '
    res += str(random.randint(0,9))
    return str(res)

def get_heure_d():
    fake = Faker()
    heure_d = fake.time('%H:%M')
    return str(heure_d)

def get_heure_d_1():
    fake = Faker()
    heure_d_1 = fake.time('%H:%M')
    return str(heure_d_1)

def get_heure_a():
    fake = Faker()
    heure_a = fake.time()
    return str(heure_a)

def get_ref_billet():
    res = ''
    for i in range(3):
        res += random.choice(string.ascii_uppercase)
    for i in range(3):
        res += str(random.randint(0,9))
    return str(res)

def get_ref_vol():
    res = 'AF'
    for i in range(3):
        res += str(random.randint(0,9))
    return str(res)

def get_pays():
    fake = Faker('fr_FR')
    return fake.country()

def get_ine():
    global ine
    ine = ''
    for j in range(10):
        ine += str(random.randint(0,9))
    ine += random.choice(string.ascii_uppercase)
    return ine

def get_ref_bourse():
    res = ''
    res += random.choice(['2020','2021','2019','2018','2017'])
    res += random.choice(['REN','BDX','MAR','NAN','STR'])
    res += ine
    return res

def get_nom_bourse():
    global nom_bourse
    fake = Faker('fr_FR')
    nom_bourse = fake.last_name()
    return nom_bourse

def get_prenom_bourse():
    global prenom_bourse
    fake = Faker('fr_FR')
    prenom_bourse = fake.first_name_male()
    return prenom_bourse

def get_email():
    res = ''
    res += nom_bourse.lower() + '.' + prenom_bourse.lower() +'@'
    res += random.choice(['gmail.com','hotmail.fr','outlook.fr','yahoo.fr','hotmail.com'])
    return res


if __name__ == '__main__':
    main()
